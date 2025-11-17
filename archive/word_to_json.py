# word_to_json.py — Conversion Word (.docx) -> data.json (HTML conservé)
# Prérequis :  pip install python-docx

from docx import Document
from pathlib import Path
import json, re, unicodedata, sys

BASE_DIR = Path("word_files")
OUT_JSON  = Path("data.json")

CATEGORY_ORDER = [
    ("Qui",       1),
    ("Quoi",      101),
    ("Ou",        201),   # tolère aussi "Ou.docx"
    ("Quand",     301),
    ("Comment",   401),
    ("Combien",   501),
    ("As",        601),
]

BRIS_PREFIX = "Bris"  # lira Bris*.docx (un ou plusieurs fichiers)

# ---------------- Utils ----------------
def normalize_filename(s: str) -> str:
    s = unicodedata.normalize("NFD", s)
    s = "".join(ch for ch in s if unicodedata.category(ch) != "Mn")
    s = s.replace("'", "")
    return s.lower()

def html_escape(s: str) -> str:
    return (s.replace("&", "&amp;").replace("<","&lt;").replace(">","&gt;"))

def run_to_html(run) -> str:
    # 1) prendre le texte brut du run
    txt = run.text.replace("\r", "")  # sécurité

    # 2) échapper d'abord <, >, &
    txt = html_escape(txt)

    # 3) convertir les retours doux APRES l'échappement
    #    (sinon <br/> deviendrait &lt;br/&gt;)
    txt = txt.replace("\n", "<br/>")

    # 4) appliquer les styles
    if getattr(run.font, "superscript", False):
        txt = f"<sup>{txt}</sup>"
    if run.italic:
        txt = f"<i>{txt}</i>"
    if run.bold:
        txt = f"<b>{txt}</b>"
    return txt

def para_to_html(p) -> str:
    # Important: conserver les paragraphes vides pour délimiter les cartes
    html = "".join(run_to_html(r) for r in p.runs)
    return html

def join_with_br(a: str, b: str) -> str:
    if not a: return b
    if not b: return a
    return a + "<br/>" + b

TAG_RE = re.compile(r"<[^>]+>")
ELLIPSIS_START = re.compile(r"^\s*…")   # « … »

def strip_tags(html: str) -> str:
    return TAG_RE.sub("", html)

# ---------------- Catégories (sans “category-aware”) ----------------
def extract_category_pairs(path: Path, cat_name: str):
    """
    Parsing générique et robuste :
      - lit le document en conservant aussi les paragraphes V I D E (délimiteurs)
      - aplatit chaque paragraphe en lignes (split sur <br/>), mais garde trace des vides
      - accumule la question jusqu'au DERNIER '?'
      - si pas de '?', pour 'Comment' on bascule en réponse dès qu'une ligne commence par '…'
      - fin de carte déclenchée par : ligne vide, ou début d'une nouvelle question (nouveau '?')
    """
    print(f"- Lecture: {path}")
    doc = Document(path)

    # 1) Aplatir en LIGNES mais on garde les paragraphes vides comme délimiteurs (None)
    lines = []
    for p in doc.paragraphs:
        html = para_to_html(p)
        if html.strip() == "":
            lines.append(None)  # délimiteur “paragraphe vide”
            continue
        # éclater les retours doux
        parts = [seg.strip() for seg in html.split("<br/>")]
        for seg in parts:
            if seg != "":
                lines.append(seg)

    pairs = []
    q_buf, a_buf = "", ""
    in_answer = False

    def flush_card():
        nonlocal q_buf, a_buf, in_answer
        if not q_buf and not a_buf:
            return
        q_text, a_text = q_buf.strip(), a_buf.strip()

        # si '?' quelque part dans q_text -> coupe proprement
        if "?" in q_text:
            cut = q_text.rfind("?")
            after = q_text[cut+1:].strip()
            q_text = q_text[:cut+1].strip()
            if after:
                a_text = join_with_br(after, a_text) if a_text else after

        pairs.append({"question": q_text, "reponse": a_text})
        q_buf, a_buf, in_answer = "", "", False

    for ln in lines:
        # délimiteur (paragraphe vide)
        if ln is None:
            if q_buf or a_buf:
                flush_card()
            continue

        ln_plain = strip_tags(ln).strip()

        if not in_answer:
            # construire la question
            q_buf = join_with_br(q_buf, ln)

            if "?" in ln_plain or "?" in q_buf:
                in_answer = True
            else:
                # Heuristique spéciale 'Comment' : la ligne qui commence par '…' est la réponse
                if cat_name == "Comment" and ELLIPSIS_START.match(ln_plain):
                    # retirer cette ligne de la question (elle appartient à la réponse)
                    # on annule l'ajout et on force en réponse
                    # (en pratique, la ligne précédente était l'énoncé sans '?')
                    # donc on enlève la dernière ligne ajoutée
                    # Pour rester simple: on bascule simplement en réponse et on met ln dedans
                    in_answer = True
                    a_buf = join_with_br(a_buf, ln)
        else:
            # phase “réponse”
            a_buf = join_with_br(a_buf, ln)

    # fin de doc -> flush
    flush_card()

    print(f"  -> {len(pairs)} Q/R")
    return pairs

# ---------------- Bris (Vrai/Faux robustes) ----------------
def extract_bris_pairs(path: Path):
    """
    Bris : affirmation + <br/> + Vrai/Faux (souvent dans le même paragraphe).
    On coupe sur <br/>, la première partie = affirmation, la suivante qui
    commence par Vrai/Faux = réponse.
    """
    print(f"- Lecture: {path}")
    doc = Document(path)

    pairs = []
    current_aff = ""

    def flush(rep_html: str = ""):
        nonlocal current_aff
        if current_aff.strip():
            pairs.append({
                "affirmation": current_aff.strip(),
                "reponse": rep_html.strip()
            })
        current_aff = ""

    for p in doc.paragraphs:
        html = para_to_html(p).strip()
        if not html:
            # paragraphe vide = séparateur éventuel
            continue

        # scinder les retours doux visibles <br/>
        segs = [s.strip() for s in html.split("<br/>") if s.strip()]
        if not segs:
            continue

        # on parcourt chaque segment ; le premier Vrai/Faux rencontré close la carte
        tmp_aff = ""
        answered = False
        for seg in segs:
            plain = strip_tags(seg).strip().lower()
            if (plain.startswith("vrai") or plain.startswith("faux")) and not answered:
                head = "Vrai" if plain.startswith("vrai") else "Faux"
                tail = strip_tags(seg)[len(head):].lstrip()
                rep_html = f"<b>{head}</b>" + ((" " + tail) if tail else "")
                # affirmation = accumulé (current_aff + tmp_aff)
                aff = current_aff
                if tmp_aff:
                    aff = join_with_br(aff, tmp_aff) if aff else tmp_aff
                current_aff = aff  # puis flush
                flush(rep_html)
                answered = True
            else:
                tmp_aff = join_with_br(tmp_aff, seg)

        # si aucun Vrai/Faux dans ce paragraphe -> ajouter à l’affirmation courante
        if not answered and tmp_aff:
            current_aff = join_with_br(current_aff, tmp_aff)

    # dernier aff orphelin
    flush()
    print(f"  -> {len(pairs)} Bris détectés")
    return pairs

# ---------------- Construction JSON ----------------
def build_json():
    out = {"categories": [], "bris": []}

    if not BASE_DIR.exists():
        print(f"ERREUR : dossier '{BASE_DIR}' introuvable")
        sys.exit(1)

    files = {normalize_filename(p.stem): p for p in BASE_DIR.glob("*.docx")}
    if not files:
        print("ERREUR : aucun .docx dans 'word_files/'")
        sys.exit(1)

    print("Docs trouvés :")
    for p in sorted(BASE_DIR.glob("*.docx")):
        print("  -", p.name)

    # Catégories
    for cat, start in CATEGORY_ORDER:
        path = files.get(normalize_filename(cat))
        if not path:
            path = files.get(normalize_filename(cat.replace("Où","Ou")))
        if not path:
            print(f"⚠ Catégorie manquante : {cat}")
            continue

        pairs = extract_category_pairs(path, cat)
        for i, pr in enumerate(pairs):
            out["categories"].append({
                "id": start + i,
                "categorie": cat,
                "question": pr["question"],
                "reponse": pr["reponse"],
            })

    # Bris : lire et fusionner tous les "Bris*.docx"
    total = 0
    for docx in sorted(BASE_DIR.glob(f"{BRIS_PREFIX}*.docx")):
        pairs = extract_bris_pairs(docx)
        for pr in pairs:
            total += 1
            out["bris"].append({
                "id": total,
                "affirmation": pr["affirmation"],
                "reponse": pr["reponse"],
            })
    print(f"→ Total Bris fusionnés : {total}")

    return out

# ---------------- Main ----------------
if __name__ == "__main__":
    print("⚙ Conversion Word → JSON…")
    data = build_json()
    OUT_JSON.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"\n✅ Écrit {OUT_JSON}  (categories={len(data['categories'])}, bris={len(data['bris'])})")
